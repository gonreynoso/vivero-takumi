#!/usr/bin/env python3
"""Genera TP_GONZALO_REYNOSO.docx para la entrega del final."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_title(doc, "Trabajo Práctico Final — Plataformas de Desarrollo")

    doc.add_paragraph("Materia: Plataformas de Desarrollo")
    doc.add_paragraph("Comisión: ACN4AV")
    doc.add_paragraph()

    add_heading(doc, "Nombre del proyecto")
    doc.add_paragraph(
        "Vivero Takumi — Tienda online y panel de gestión para un vivero de plantas "
        "(frontend React + backend Express)."
    )

    add_heading(doc, "Integrantes del grupo")
    add_table(doc, ["Nombre", "Email"], [["Gonzalo Reynoso", "gonzalo.reynoso@davinci.edu.ar"]])

    add_heading(doc, "Temática")
    doc.add_paragraph(
        "Vivero de plantas orientado a venta minorista y cuidado post-compra. "
        "La empresa es ficticia: un vivero local que digitaliza catálogo, pedidos "
        "y gestión interna mediante una SPA y una API REST."
    )

    add_heading(doc, "Arquitectura del sistema")
    doc.add_paragraph(
        "El proyecto cumple la consigna del final con dos aplicaciones separadas:"
    )
    doc.add_paragraph("Frontend (frontend/): SPA React consumida en http://localhost:5173", style="List Bullet")
    doc.add_paragraph("Backend (backend/): API REST Express + MySQL en http://localhost:8888", style="List Bullet")
    doc.add_paragraph(
        "El frontend no accede a la base de datos directamente; obtiene y persiste datos "
        "mediante fetch al API (Authorization: Bearer <JWT>)."
    )

    add_heading(doc, "Requisitos del final (cumplidos)")
    add_table(
        doc,
        ["Requisito", "Implementación"],
        [
            ["Mínimo 2 tipos de usuarios", "4 roles: admin, manager, empleado, cliente"],
            ["Seguridad con access token", "JWT + bcrypt; middleware requireAuth y requireRole"],
            ["API REST", "Express en backend/ — auth, plantas, categorías, pedidos, usuarios"],
            ["Frontend y backend separados", "Carpetas frontend/ y backend/; dos procesos y puertos"],
            ["Persistencia en servidor", "MySQL 8 con esquema en backend/sql/schema.sql"],
        ],
    )

    add_heading(doc, "Estructura del proyecto")
    doc.add_paragraph(
        "vivero-takumi/\n"
        "├── frontend/          # React + Vite\n"
        "│   ├── src/\n"
        "│   ├── public/\n"
        "│   └── .env           # VITE_API_URL\n"
        "├── backend/           # Express + MySQL\n"
        "│   ├── app.js\n"
        "│   ├── sql/\n"
        "│   └── src/           # routes, controllers, models, middlewares\n"
        "└── TP_GONZALO_REYNOSO.docx"
    )

    add_heading(doc, "Usuarios y roles")
    doc.add_paragraph("La aplicación define cuatro roles con permisos distintos:")
    add_table(
        doc,
        ["Rol", "Descripción", "Acceso principal"],
        [
            ["admin", "Administrador del sistema", "Dashboard, plantas, categorías, pedidos, usuarios"],
            ["manager", "Encargado del vivero", "Dashboard, plantas, categorías, pedidos"],
            ["empleado", "Personal de depósito", "Stock y pedidos (actualizar estados)"],
            ["cliente", "Comprador registrado", "Catálogo, carrito, checkout, historial y riego"],
        ],
    )

    add_heading(doc, "Usuarios de prueba")
    add_table(
        doc,
        ["Rol", "Email", "Contraseña"],
        [
            ["Admin", "gonzalo.reynoso9@gmail.com", "Ver backend/scripts/seed.js"],
            ["Manager", "manager@viverotakumi.com", "manager123"],
            ["Empleado", "empleado@viverotakumi.com", "emp123"],
            ["Cliente", "cliente@viverotakumi.com", "cli123"],
        ],
    )

    add_heading(doc, "Funcionalidades")
    add_heading(doc, "Requisitos de la consigna", level=2)
    for item in [
        "Single Page Application (React + Vite + React Router)",
        "API REST con Express, MySQL y arquitectura MVC (routes → controllers → models)",
        "Autenticación con JWT (access token 15 min + refresh token 7 días) y bcrypt en contraseñas",
        "Login, registro, logout y sesión persistente en el navegador",
        "Gestión de usuarios: alta, edición, baja y asignación de rol (panel admin)",
        "Cuatro roles de usuario diferenciados",
        "Gestión de información mediante interfaz gráfica (CRUD de plantas, categorías, pedidos y usuarios)",
        "Componentes visuales separados y reutilizables",
        "Frontend y backend desacoplados; CORS configurado para localhost:5173",
        "Documentación OpenAPI/Swagger en /api/docs",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "Funcionalidades adicionales (originalidad)", level=2)
    doc.add_paragraph("Tienda (cliente / invitado):", style="List Bullet")
    for item in [
        "Home con hero, categorías, destacados, beneficios, testimonios y FAQ",
        "Catálogo con búsqueda y filtros",
        "Detalle de planta con galería y guía de cuidado",
        "Recomendador de plantas (quiz interactivo)",
        "Carrito y checkout (envío a domicilio o retiro en local)",
        "Registro de clientes y vista Mis compras (pedidos + seguimiento de riego)",
        "Modo oscuro y diseño responsive",
    ]:
        doc.add_paragraph(item, style="List Bullet 2")

    doc.add_paragraph("Panel administrativo:", style="List Bullet")
    for item in [
        "Dashboard con KPIs y gráficos",
        "CRUD de plantas, categorías y pedidos",
        "Gestión de usuarios (solo admin)",
    ]:
        doc.add_paragraph(item, style="List Bullet 2")

    doc.add_paragraph("Panel empleado:", style="List Bullet")
    doc.add_paragraph("Vista y actualización de stock", style="List Bullet 2")
    doc.add_paragraph("Gestión de pedidos (cambio de estado)", style="List Bullet 2")

    add_heading(doc, "Endpoints principales del API")
    add_table(
        doc,
        ["Método", "Ruta", "Auth", "Descripción"],
        [
            ["POST", "/api/auth/register", "No", "Registro de cliente"],
            ["POST", "/api/auth/login", "No", "Login → accessToken + refreshToken"],
            ["GET", "/api/auth/refresh-token", "Refresh", "Renovar access token"],
            ["GET", "/api/auth/me", "Sí", "Usuario actual"],
            ["GET", "/api/plantas", "Opcional", "Catálogo público"],
            ["GET/POST/PUT/DELETE", "/api/plantas/:ID", "Staff", "CRUD plantas"],
            ["GET", "/api/categorias", "No", "Listado de categorías"],
            ["POST/PUT/DELETE", "/api/categorias", "Staff", "CRUD categorías"],
            ["GET", "/api/pedidos", "Sí", "Pedidos (cliente ve solo los suyos)"],
            ["POST", "/api/pedidos", "No*", "Crear pedido (checkout invitado o logueado)"],
            ["PATCH", "/api/pedidos/:ID/estado", "Staff", "Cambiar estado del pedido"],
            ["GET/POST/PUT/DELETE", "/api/usuarios", "Admin", "Gestión de usuarios"],
        ],
    )
    doc.add_paragraph("* El checkout funciona con o sin sesión; las rutas de staff exigen JWT.")

    add_heading(doc, "Tecnologías utilizadas")
    add_heading(doc, "Frontend", level=2)
    add_table(
        doc,
        ["Capa", "Tecnología"],
        [
            ["Framework", "React 19"],
            ["Build", "Vite 8"],
            ["Routing", "React Router 7"],
            ["Estilos", "Tailwind CSS 3"],
            ["Gráficos", "Recharts"],
            ["Cliente HTTP", "fetch nativo (frontend/src/lib/api.js)"],
        ],
    )

    add_heading(doc, "Backend", level=2)
    add_table(
        doc,
        ["Capa", "Tecnología"],
        [
            ["Runtime", "Node.js 18+"],
            ["Framework", "Express 5"],
            ["Base de datos", "MySQL 8 (mysql2)"],
            ["Autenticación", "jsonwebtoken + bcrypt"],
            ["CORS", "cors"],
            ["Documentación", "swagger-ui-express (/api/docs)"],
        ],
    )

    add_heading(doc, "Ejecución del proyecto")
    add_heading(doc, "1. Base de datos (una sola vez)", level=2)
    doc.add_paragraph("sudo mysql < backend/sql/setup-local.sql")
    doc.add_paragraph("cd backend && cp .env.example .env && npm install && npm run seed")

    add_heading(doc, "2. Backend", level=2)
    doc.add_paragraph("cd backend && npm run dev")
    doc.add_paragraph("API: http://localhost:8888 — Swagger: http://localhost:8888/api/docs")

    add_heading(doc, "3. Frontend", level=2)
    doc.add_paragraph("cd frontend && cp .env.example .env && pnpm install && pnpm dev")
    doc.add_paragraph("SPA: http://localhost:5173")

    add_heading(doc, "Demo sugerida para la defensa")
    for i, item in enumerate(
        [
            "Catálogo sin login — datos desde GET /api/plantas",
            "Registro y login de cliente — JWT en localStorage",
            "Checkout — POST /api/pedidos y descuento de stock en MySQL",
            "Mis compras — GET /api/pedidos filtrado por usuario",
            "Login admin — CRUD plantas y usuarios con token",
            "Login empleado — actualizar stock y estado de pedidos",
            "Postman — ruta protegida sin token (401) y con token (200)",
        ],
        start=1,
    ):
        doc.add_paragraph(f"{i}. {item}")

    add_heading(doc, "Repositorio y entrega")
    doc.add_paragraph("Código fuente: https://github.com/gonreynoso/vivero-takumi")
    doc.add_paragraph()
    doc.add_paragraph("Formato del ZIP de entrega:")
    doc.add_paragraph(
        "TP_GONZALO_REYNOSO.zip\n"
        "├── TP_GONZALO_REYNOSO.docx\n"
        "├── frontend/\n"
        "└── backend/"
    )
    doc.add_paragraph(
        "Comando sugerido:\n"
        "zip -r TP_GONZALO_REYNOSO.zip TP_GONZALO_REYNOSO.docx frontend backend "
        '-x "frontend/node_modules/*" -x "backend/node_modules/*" -x "frontend/dist/*"'
    )

    return doc


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    output = root / "TP_GONZALO_REYNOSO.docx"
    build_document().save(output)
    print(f"Generado: {output}")


if __name__ == "__main__":
    main()
